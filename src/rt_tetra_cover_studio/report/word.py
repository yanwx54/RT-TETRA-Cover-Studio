from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from .common import build_curve_chart_images, project_info_rows
from ..paths import resource_path


DEFAULT_TEMPLATE_PATH = resource_path("config", "report_template.json")


def export_word_report(
    report_data: dict[str, Any],
    output_path: str | Path,
    template_path: str | Path = DEFAULT_TEMPLATE_PATH,
) -> Path:
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError as exc:  # pragma: no cover - exercised only without dependency.
        raise RuntimeError("python-docx is required to export Word reports.") from exc

    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    template = _load_template(template_path)

    document = Document()
    _apply_header_footer(document, template)
    document.add_heading(template["title"], level=0)
    document.add_paragraph(template["subtitle"])

    _add_project_info(document, project_info_rows(template, report_data))
    _add_summary(document, report_data["summary"])
    _add_input_table(document, report_data["input"])
    _add_calculation_sections(document, report_data["details"]["calculation_sections"])
    _add_iteration_summary(document, report_data["details"]["iteration_steps"])
    _add_curve_summary(document, report_data["charts"], template["curve_sample_points"])
    _add_curve_images(document, report_data["charts"], Inches)
    _add_warnings(document, report_data["summary"].get("warnings", []))

    document.save(path)
    return path


def _load_template(template_path: str | Path) -> dict[str, Any]:
    with Path(template_path).open("r", encoding="utf-8") as file:
        return json.load(file)


def _apply_header_footer(document: Any, template: dict[str, Any]) -> None:
    section = document.sections[0]
    section.header.paragraphs[0].text = template["title"]
    section.footer.paragraphs[0].text = template.get("footer", "")


def _add_project_info(document: Any, rows: list[tuple[str, str]]) -> None:
    document.add_heading("工程信息", level=1)
    _add_key_value_table(document, rows)


def _add_summary(document: Any, summary: dict[str, Any]) -> None:
    document.add_heading("核心结论", level=1)
    rows = [
        ("传播模型", summary["model_name"]),
        ("最大覆盖距离", f"{summary['coverage_distance_m']:.1f} m"),
        ("受限链路", summary["limiting_link"]),
        ("模型校准状态", summary["calibration_status"]),
        ("模型参数来源", summary["calibration_source"] or "未提供"),
        ("基站 EIRP", f"{summary['base_eirp_dbm']:.2f} dBm"),
        ("手台 EIRP", f"{summary['mobile_eirp_dbm']:.2f} dBm"),
        ("下行 MAPL", f"{summary['downlink_mapl_db']:.2f} dB"),
        ("上行 MAPL", f"{summary['uplink_mapl_db']:.2f} dB"),
        ("系统 MAPL", f"{summary['max_path_loss_db']:.2f} dB"),
        ("下行边界 RSSI", f"{summary['boundary_rssi_dbm']:.2f} dBm"),
    ]
    _add_key_value_table(document, rows)


def _add_input_table(document: Any, input_data: dict[str, Any]) -> None:
    document.add_heading("输入参数", level=1)
    rows = []
    for key, value in input_data.items():
        if key == "scenario_params":
            for param_key, param_value in value.items():
                rows.append((f"scenario_params.{param_key}", param_value))
        else:
            rows.append((key, value))
    _add_key_value_table(document, rows)


def _add_calculation_sections(document: Any, sections: list[dict[str, Any]]) -> None:
    document.add_heading("详细计算过程", level=1)
    for section in sections:
        document.add_heading(f"{section['number']} {section['title']}", level=2)
        document.add_paragraph(section["description"])
        for detail in section["details"]:
            document.add_paragraph(detail["name"], style="Heading 3")
            document.add_paragraph(f"公式：{detail['formula']}")
            document.add_paragraph(f"代入：{detail['substitution']}")
            document.add_paragraph(f"结果：{detail['result']}")


def _add_iteration_summary(document: Any, iteration_steps: list[dict[str, Any]]) -> None:
    document.add_heading("覆盖距离求解", level=1)
    document.add_paragraph(f"迭代记录数量：{len(iteration_steps)}")
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    _set_cells(table.rows[0].cells, ["迭代", "距离(m)", "Path Loss(dB)"])
    for step in iteration_steps[:10]:
        cells = table.add_row().cells
        _set_cells(
            cells,
            [
                str(step["iteration"]),
                f"{step['distance_m']:.1f}",
                f"{step['path_loss_db']:.2f}",
            ],
        )


def _add_curve_summary(
    document: Any, charts: dict[str, Any], sample_points: int
) -> None:
    document.add_heading("曲线数据摘要", level=1)
    boundary = charts["coverage_boundary"]
    document.add_paragraph(
        "覆盖边界："
        f"{boundary['distance_m']:.1f} m，"
        f"Path Loss {boundary['path_loss_db']:.2f} dB，"
        f"RSSI {boundary['rssi_dbm']:.2f} dBm"
    )

    path_loss_curve = charts["path_loss_curve"]
    rssi_curve = charts["rssi_curve"]
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    _set_cells(table.rows[0].cells, ["距离(m)", "Path Loss(dB)", "RSSI(dBm)"])
    for path_point, rssi_point in list(zip(path_loss_curve, rssi_curve))[:sample_points]:
        cells = table.add_row().cells
        _set_cells(
            cells,
            [
                f"{path_point['distance_m']:.1f}",
                f"{path_point['path_loss_db']:.2f}",
                f"{rssi_point['rssi_dbm']:.2f}",
            ],
        )


def _add_curve_images(document: Any, charts: dict[str, Any], inches: Any) -> None:
    document.add_heading("覆盖曲线图", level=1)
    for chart in build_curve_chart_images(charts):
        document.add_paragraph(chart.title)
        chart.image.seek(0)
        document.add_picture(chart.image, width=inches(6.2))


def _add_warnings(document: Any, warnings: list[str]) -> None:
    document.add_heading("计算提示", level=1)
    if warnings:
        for warning in warnings:
            document.add_paragraph(warning, style="List Bullet")
    else:
        document.add_paragraph("无告警")


def _add_key_value_table(document: Any, rows: list[tuple[str, Any]]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    _set_cells(table.rows[0].cells, ["项目", "值"])
    for key, value in rows:
        cells = table.add_row().cells
        _set_cells(cells, [str(key), str(value)])


def _set_cells(cells: Any, values: list[str]) -> None:
    for cell, value in zip(cells, values):
        cell.text = value
