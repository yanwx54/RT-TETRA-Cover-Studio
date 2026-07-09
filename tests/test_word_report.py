from __future__ import annotations

import sys
import unittest
from pathlib import Path


PROJECT_DIR = Path(__file__).resolve().parents[1]
SRC_DIR = PROJECT_DIR / "src"
if str(SRC_DIR) not in sys.path:
    sys.path.insert(0, str(SRC_DIR))

from docx import Document

from rt_tetra_cover_studio.engine import calculate_coverage
from rt_tetra_cover_studio.io import calculation_result_to_dict, load_example_case
from rt_tetra_cover_studio.report import export_word_report


class WordReportTest(unittest.TestCase):
    def test_export_word_report_from_standard_example(self) -> None:
        case_data = load_example_case(PROJECT_DIR / "examples" / "underground_standard.json")
        result = calculate_coverage(case_data["calculation_input"])
        report_data = calculation_result_to_dict(result)
        output_path = PROJECT_DIR / "reports" / "test_outputs" / "word_report.docx"

        saved_path = export_word_report(report_data, output_path)

        self.assertTrue(saved_path.exists())
        self.assertGreater(saved_path.stat().st_size, 0)
        document = Document(saved_path)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("RT-TETRA Cover Studio 覆盖计算报告", text)
        self.assertIn("工程信息", text)
        self.assertIn("核心结论", text)
        self.assertGreaterEqual(len(document.tables), 5)
        self.assertGreaterEqual(len(document.inline_shapes), 2)
        header_text = "\n".join(
            paragraph.text for section in document.sections for paragraph in section.header.paragraphs
        )
        footer_text = "\n".join(
            paragraph.text for section in document.sections for paragraph in section.footer.paragraphs
        )
        self.assertIn("RT-TETRA Cover Studio 覆盖计算报告", header_text)
        self.assertIn("RT-TETRA Cover Studio", footer_text)


if __name__ == "__main__":
    unittest.main()
